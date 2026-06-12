from __future__ import annotations

import re
from io import BytesIO

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

_HEADING_RE = re.compile(r'^(#{1,3})\s+(.+?)\s*$')
_BULLET_RE = re.compile(r'^[-*]\s+(.+?)\s*$')
_NUMBERED_RE = re.compile(r'^\d+\.\s+(.+?)\s*$')
_BOLD_RE = re.compile(r'\*\*(.+?)\*\*')


# def export_summary_docx để xuất summary docx.
# vd: nhận đầu vào -> trả kết quả đã xử lý.
def export_summary_docx(summary) -> bytes:
    buffer = BytesIO()
    doc = DocxDocument()

    title = _summary_title(summary)
    doc.add_heading(f'Tóm tắt: {title}', level=1)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f'Ngày tóm tắt: {_summary_created_at(summary)}\n'
        f'Người tóm tắt: {_summary_created_by(summary)}\n'
        f'Mô hình AI: {_summary_model_name(summary)}'
    )
    meta_run.italic = True

    doc.add_paragraph()
    _render_markdown_to_docx(doc, _summary_text(summary))
    doc.add_paragraph()

    footer = doc.add_paragraph('— Hết —')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(buffer)
    return buffer.getvalue()


# def export_summary_md để xuất summary md.
# vd: nhận đầu vào -> trả kết quả đã xử lý.
def export_summary_md(summary) -> str:
    lines = [
        '---',
        f'title: "{_yaml_escape(f"Tóm tắt: {_summary_title(summary)}")}"',
        f'created_at: "{_summary_created_at_iso(summary)}"',
        f'created_by: "{_yaml_escape(_summary_created_by(summary))}"',
        f'model: "{_yaml_escape(_summary_model_name(summary))}"',
        '---',
        '',
        _summary_text(summary).strip(),
        '',
    ]
    return '\n'.join(lines)


# def _render_markdown_to_docx để kết xuất markdown to docx.
# vd: nhận tham số đầu vào -> trả cấu trúc dữ liệu/chuỗi đã dựng.
def _render_markdown_to_docx(doc: DocxDocument, text: str) -> None:
    in_code_block = False
    for raw_line in str(text or '').splitlines():
        line = raw_line.rstrip()
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            paragraph = doc.add_paragraph(line)
            paragraph.style = 'Intense Quote'
            for run in paragraph.runs:
                run.font.name = 'Consolas'
            continue
        if not line.strip():
            doc.add_paragraph()
            continue
        heading_match = _HEADING_RE.match(line)
        if heading_match:
            doc.add_heading(
                heading_match.group(2).strip(),
                level=min(len(heading_match.group(1)), 3),
            )
            continue
        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            _write_inline_runs(
                doc.add_paragraph(style='List Bullet'),
                bullet_match.group(1).strip(),
            )
            continue
        numbered_match = _NUMBERED_RE.match(line)
        if numbered_match:
            _write_inline_runs(
                doc.add_paragraph(style='List Number'),
                numbered_match.group(1).strip(),
            )
            continue
        _write_inline_runs(doc.add_paragraph(), line)


# def _write_inline_runs để write inline runs.
# vd: nhận đầu vào -> trả kết quả đã xử lý.
def _write_inline_runs(paragraph, text: str) -> None:
    last_index = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > last_index:
            paragraph.add_run(text[last_index:match.start()])
        paragraph.add_run(match.group(1)).bold = True
        last_index = match.end()
    if last_index < len(text):
        paragraph.add_run(text[last_index:])


# def _summary_title để summary title.
# vd: nhận đầu vào -> trả kết quả đã xử lý.
def _summary_title(summary) -> str:
    document = getattr(summary, 'document', None)
    return str(getattr(document, 'title', None) or 'Văn bản')


# def _summary_text để summary text.
# vd: nhận đầu vào -> trả kết quả đã xử lý.
def _summary_text(summary) -> str:
    for attr in ('content_md', 'summary', 'content_text'):
        value = getattr(summary, attr, None)
        if str(value or '').strip():
            return str(value).strip()
    return ''


# def _summary_created_by để summary created by.
# vd: nhận đầu vào -> trả kết quả đã xử lý.
def _summary_created_by(summary) -> str:
    explicit_name = str(getattr(summary, 'created_by_name', '') or '').strip()
    if explicit_name:
        return explicit_name
    user = getattr(summary, 'created_by', None)
    if user is None:
        return 'Hệ thống'
    return getattr(user, 'get_full_name', lambda: '')().strip() or getattr(
        user,
        'username',
        'Hệ thống',
    )


# def _summary_model_name để summary model name.
# vd: nhận đầu vào -> trả kết quả đã xử lý.
def _summary_model_name(summary) -> str:
    return str(getattr(summary, 'model_name', None) or '—')


# def _summary_created_at để summary created at.
# vd: nhận đầu vào -> trả kết quả đã xử lý.
def _summary_created_at(summary) -> str:
    created_at = getattr(summary, 'created_at', None)
    if created_at is None:
        return '—'
    return created_at.strftime('%d/%m/%Y %H:%M')


# def _summary_created_at_iso để summary created at iso.
# vd: nhận đầu vào -> trả kết quả đã xử lý.
def _summary_created_at_iso(summary) -> str:
    created_at = getattr(summary, 'created_at', None)
    if created_at is None:
        return ''
    return created_at.isoformat()


# def _yaml_escape để yaml escape.
# vd: nhận đầu vào -> trả kết quả đã xử lý.
def _yaml_escape(value: str) -> str:
    return str(value or '').replace('\\', '\\\\').replace('"', '\\"')

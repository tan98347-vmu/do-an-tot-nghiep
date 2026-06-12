"""
Thuoc chuc nang nao: Tao mau van ban, Mau dung chung, Mau phong ban, Mau rieng, Mau yeu thich va Tat ca mau (Admin).
Vai tro backend: File `document_templates/utils.py` giu hoac ho tro luong backend cho CRUD mau, duyet mau, version mau, import DOCX/URL, bulk upload va preview noi dung mau.
Vai tro cua no trong frontend: Cac man `/templates`, `/templates/create`, man chi tiet mau va man bulk upload lay du lieu hoac chiu tac dong gian tiep tu file nay.
Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `api/serializers/templates.py`, `accounts.permissions`, `document_templates.models`, `document_templates.utils`, `document_templates.versioning`.
Tac dung: Giu cho du lieu mau, quyen thao tac va preview cua nhom man Mau van ban luon dong nhat giua API, storage va chi so tim kiem.
"""

import io
import re
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Cu phap bien dung chung: `{{ten_bien}}`, cho phep co khoang trang quanh ten
# (vi du nguoi dung go thu cong `{{ ten_bien }}`). Ten bien la \w+ (ho tro Unicode).
VARIABLE_TOKEN_RE = re.compile(r'\{\{\s*(\w+)\s*\}\}')


# def extract_template_variables trích tập tên biến {{ten}} từ nội dung (bỏ HTML, cho phép khoảng trắng quanh tên), trả danh sách đã sắp xếp + khử trùng.
# vd: '{{ho_ten}} {{ ngay }}' -> ['ho_ten','ngay'].
def extract_template_variables(content):
    """Trich xuat danh sach ten bien duy nhat tu noi dung mau (da bo tag HTML).

    Cho phep khoang trang quanh ten bien nen ca `{{ten}}` lan `{{ ten }}` deu nhan.
    """
    plain = re.sub(r'<[^>]+>', '', content or '')
    return sorted({match for match in VARIABLE_TOKEN_RE.findall(plain) if match})


# def fill_variables_in_text thay mọi {{ten}} trong text bằng giá trị tương ứng (cho phép khoảng trắng quanh tên).
# vd: '{{ho_ten}}' + {'ho_ten':'A'} -> 'A'.
def fill_variables_in_text(text, variables):
    """Thay the cac placeholder `{{ten}}`/`{{ ten }}` trong text bang gia tri."""
    if not text:
        return text
    result = text
    for key, value in (variables or {}).items():
        pattern = re.compile(r'\{\{\s*' + re.escape(str(key)) + r'\s*\}\}')
        result = pattern.sub(lambda _m, _v=str(value): _v, result)
    return result

# def extract_text_from_docx đọc toàn bộ text từ file DOCX (đoạn văn + ô bảng) để dò biến / đánh chỉ mục.
# vd: file DOCX -> chuỗi text gồm các dòng nội dung.
def extract_text_from_docx(docx_file):
    """
    Thuoc chuc nang nao: Tao mau van ban, Mau dung chung, Mau phong ban, Mau rieng, Mau yeu thich va Tat ca mau (Admin).
    Vai tro backend: Ham `extract_text_from_docx` la mot don vi xu ly backend cua file `document_templates/utils.py`, chu yeu de trich xuat noi dung hoac gia tri trung gian.
    Vai tro cua no trong frontend: Frontend chu yeu quan sat ket qua gian tiep cua buoc trich xuat noi dung hoac gia tri trung gian nay thong qua API, du lieu luu tru hoac trang thai do lop goi phia tren tra ve.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `api/serializers/templates.py`, `accounts.permissions`, `document_templates.models`, `document_templates.utils`, `document_templates.versioning`. Dung cung cap voi cac ham `docx_preview_as_html`, `_replace_in_paragraph`, `_replace_in_table` trong module nay.
    Tac dung: Tach rieng trach nhiem trich xuat noi dung hoac gia tri trung gian de pham vi tac dong cua `extract_text_from_docx` ro rang hon.
    """
    doc = DocxDocument(docx_file)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        lines.append(para.text)
    return '\n'.join(lines)

# def docx_preview_as_html chuyển file DOCX sang HTML (mammoth) để xem trước trên web; lỗi -> None.
# vd: mẫu DOCX -> HTML preview hiển thị trong trình duyệt.
def docx_preview_as_html(docx_path):
    """
    Thuoc chuc nang nao: Tao mau van ban, Mau dung chung, Mau phong ban, Mau rieng, Mau yeu thich va Tat ca mau (Admin).
    Vai tro backend: Ham `docx_preview_as_html` la mot don vi xu ly backend cua file `document_templates/utils.py`, chu yeu de chuan bi noi dung xem truoc.
    Vai tro cua no trong frontend: Frontend chu yeu quan sat ket qua gian tiep cua buoc chuan bi noi dung xem truoc nay thong qua API, du lieu luu tru hoac trang thai do lop goi phia tren tra ve.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `api/serializers/templates.py`, `accounts.permissions`, `document_templates.models`, `document_templates.utils`, `document_templates.versioning`. Dung cung cap voi cac ham `extract_text_from_docx`, `_replace_in_paragraph`, `_replace_in_table` trong module nay.
    Tac dung: Tach rieng trach nhiem chuan bi noi dung xem truoc de pham vi tac dong cua `docx_preview_as_html` ro rang hon.
    """
    try:
        import mammoth
        with open(docx_path, 'rb') as f:
            result = mammoth.convert_to_html(f)
        return result.value
    except Exception:
        return None

# def _replace_in_paragraph thay biến trong 1 đoạn văn DOCX nhưng GIỮ định dạng: gộp text các run, điền biến, rồi đặt lại vào run đầu và xóa text các run sau.
# vd: đoạn '{{ho_ten}}' -> 'A' mà vẫn giữ font/format.
def _replace_in_paragraph(para, variables):
    """
    Thuoc chuc nang nao: Tao mau van ban, Mau dung chung, Mau phong ban, Mau rieng, Mau yeu thich va Tat ca mau (Admin).
    Vai tro backend: Ham `_replace_in_paragraph` la mot don vi xu ly backend cua file `document_templates/utils.py`, chu yeu de thuc hien phan xu ly chuyen trach cua symbol hien tai.
    Vai tro cua no trong frontend: Frontend chu yeu quan sat ket qua gian tiep cua buoc thuc hien phan xu ly chuyen trach cua symbol hien tai nay thong qua API, du lieu luu tru hoac trang thai do lop goi phia tren tra ve.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `api/serializers/templates.py`, `accounts.permissions`, `document_templates.models`, `document_templates.utils`, `document_templates.versioning`. Thuong duoc cac ham public nhu `extract_text_from_docx`, `docx_preview_as_html`, `render_docx_from_template` goi lai.
    Tac dung: Tach rieng trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai de pham vi tac dong cua `_replace_in_paragraph` ro rang hon.
    """
    full_text = ''.join(run.text for run in para.runs)
    if not full_text:
        return

    new_text = fill_variables_in_text(full_text, variables)

    if new_text == full_text:
        return

    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''

# def _replace_in_table thay biến trong mọi ô của 1 bảng DOCX.
# vd: bảng có ô '{{ngay}}' -> điền giá trị vào ô.
def _replace_in_table(table, variables):
    """
    Thuoc chuc nang nao: Tao mau van ban, Mau dung chung, Mau phong ban, Mau rieng, Mau yeu thich va Tat ca mau (Admin).
    Vai tro backend: Ham `_replace_in_table` la mot don vi xu ly backend cua file `document_templates/utils.py`, chu yeu de thuc hien phan xu ly chuyen trach cua symbol hien tai.
    Vai tro cua no trong frontend: Frontend chu yeu quan sat ket qua gian tiep cua buoc thuc hien phan xu ly chuyen trach cua symbol hien tai nay thong qua API, du lieu luu tru hoac trang thai do lop goi phia tren tra ve.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `api/serializers/templates.py`, `accounts.permissions`, `document_templates.models`, `document_templates.utils`, `document_templates.versioning`. Thuong duoc cac ham public nhu `extract_text_from_docx`, `docx_preview_as_html`, `render_docx_from_template` goi lai.
    Tac dung: Tach rieng trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai de pham vi tac dong cua `_replace_in_table` ro rang hon.
    """
    for row in table.rows:
        for cell in row.cells:
            _replace_in_cell(cell, variables)

# def _replace_in_cell thay biến trong 1 ô bảng (cả đoạn văn lẫn bảng lồng bên trong).
# vd: ô chứa bảng con -> đệ quy thay biến.
def _replace_in_cell(cell, variables):
    """
    Thuoc chuc nang nao: Tao mau van ban, Mau dung chung, Mau phong ban, Mau rieng, Mau yeu thich va Tat ca mau (Admin).
    Vai tro backend: Ham `_replace_in_cell` la mot don vi xu ly backend cua file `document_templates/utils.py`, chu yeu de thuc hien phan xu ly chuyen trach cua symbol hien tai.
    Vai tro cua no trong frontend: Frontend chu yeu quan sat ket qua gian tiep cua buoc thuc hien phan xu ly chuyen trach cua symbol hien tai nay thong qua API, du lieu luu tru hoac trang thai do lop goi phia tren tra ve.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `api/serializers/templates.py`, `accounts.permissions`, `document_templates.models`, `document_templates.utils`, `document_templates.versioning`. Thuong duoc cac ham public nhu `extract_text_from_docx`, `docx_preview_as_html`, `render_docx_from_template` goi lai.
    Tac dung: Tach rieng trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai de pham vi tac dong cua `_replace_in_cell` ro rang hon.
    """
    for para in cell.paragraphs:
        _replace_in_paragraph(para, variables)
    for nested_table in cell.tables:
        _replace_in_table(nested_table, variables)

# def _replace_in_document_part thay biến trong 1 phần tài liệu (thân / header / footer): mọi đoạn văn + bảng.
# vd: dùng cho cả phần thân lẫn header/footer của DOCX.
def _replace_in_document_part(doc_part, variables):
    """
    Thuoc chuc nang nao: Tao mau van ban, Mau dung chung, Mau phong ban, Mau rieng, Mau yeu thich va Tat ca mau (Admin).
    Vai tro backend: Ham `_replace_in_document_part` la mot don vi xu ly backend cua file `document_templates/utils.py`, chu yeu de xu ly du lieu hoac thao tac lien quan toi van ban.
    Vai tro cua no trong frontend: Frontend chu yeu quan sat ket qua gian tiep cua buoc xu ly du lieu hoac thao tac lien quan toi van ban nay thong qua API, du lieu luu tru hoac trang thai do lop goi phia tren tra ve.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `api/serializers/templates.py`, `accounts.permissions`, `document_templates.models`, `document_templates.utils`, `document_templates.versioning`. Thuong duoc cac ham public nhu `extract_text_from_docx`, `docx_preview_as_html`, `render_docx_from_template` goi lai.
    Tac dung: Tach rieng trach nhiem xu ly du lieu hoac thao tac lien quan toi van ban de pham vi tac dong cua `_replace_in_document_part` ro rang hon.
    """
    for para in getattr(doc_part, 'paragraphs', []):
        _replace_in_paragraph(para, variables)
    for table in getattr(doc_part, 'tables', []):
        _replace_in_table(table, variables)

# def render_docx_from_template mở file DOCX mẫu, thay biến trong thân + header + footer (giữ nguyên định dạng), trả file DOCX (BytesIO) đã điền.
# vd: mẫu DOCX 'Đơn xin nghỉ' + biến -> file DOCX điền sẵn, giữ layout.
def render_docx_from_template(docx_path, variables):
    """
    Thuoc chuc nang nao: Tao mau van ban, Mau dung chung, Mau phong ban, Mau rieng, Mau yeu thich va Tat ca mau (Admin).
    Vai tro backend: Ham `render_docx_from_template` la mot don vi xu ly backend cua file `document_templates/utils.py`, chu yeu de render noi dung dau ra de tra ve hoac luu tru.
    Vai tro cua no trong frontend: Frontend chu yeu quan sat ket qua gian tiep cua buoc render noi dung dau ra de tra ve hoac luu tru nay thong qua API, du lieu luu tru hoac trang thai do lop goi phia tren tra ve.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `api/serializers/templates.py`, `accounts.permissions`, `document_templates.models`, `document_templates.utils`, `document_templates.versioning`. Dung cung cap voi cac ham `extract_text_from_docx`, `docx_preview_as_html`, `_replace_in_paragraph` trong module nay.
    Tac dung: Tach rieng trach nhiem render noi dung dau ra de tra ve hoac luu tru de pham vi tac dong cua `render_docx_from_template` ro rang hon.
    """
    doc = DocxDocument(docx_path)

    _replace_in_document_part(doc, variables)

    for section in doc.sections:
        _replace_in_document_part(section.header, variables)
        _replace_in_document_part(section.footer, variables)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# def create_docx_from_html dựng file DOCX từ nội dung HTML (html2docx); HTML rỗng -> 1 đoạn trống.
# vd: mẫu thủ công (content HTML) -> file DOCX.
def create_docx_from_html(html_content):
    """
    Thuoc chuc nang nao: Tao mau van ban, Mau dung chung, Mau phong ban, Mau rieng, Mau yeu thich va Tat ca mau (Admin).
    Vai tro backend: Ham `create_docx_from_html` la mot don vi xu ly backend cua file `document_templates/utils.py`, chu yeu de tao moi ban ghi hoac khoi tao mot luong xu ly.
    Vai tro cua no trong frontend: Frontend chu yeu quan sat ket qua gian tiep cua buoc tao moi ban ghi hoac khoi tao mot luong xu ly nay thong qua API, du lieu luu tru hoac trang thai do lop goi phia tren tra ve.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `api/serializers/templates.py`, `accounts.permissions`, `document_templates.models`, `document_templates.utils`, `document_templates.versioning`. Dung cung cap voi cac ham `extract_text_from_docx`, `docx_preview_as_html`, `_replace_in_paragraph` trong module nay.
    Tac dung: Tach rieng trach nhiem tao moi ban ghi hoac khoi tao mot luong xu ly de pham vi tac dong cua `create_docx_from_html` ro rang hon.
    """
    from html2docx import html2docx as _html2docx

    
    
    if not html_content or not html_content.strip():
        html_content = '<p></p>'

    buf = _html2docx(html_content, title='Document')
    buf.seek(0)
    return buf

# def create_docx_from_text dựng file DOCX đơn giản từ text thuần (mỗi dòng 1 đoạn).
# vd: text nhiều dòng -> DOCX mỗi dòng 1 paragraph.
def create_docx_from_text(text):
    """
    Thuoc chuc nang nao: Tao mau van ban, Mau dung chung, Mau phong ban, Mau rieng, Mau yeu thich va Tat ca mau (Admin).
    Vai tro backend: Ham `create_docx_from_text` la mot don vi xu ly backend cua file `document_templates/utils.py`, chu yeu de tao moi ban ghi hoac khoi tao mot luong xu ly.
    Vai tro cua no trong frontend: Frontend chu yeu quan sat ket qua gian tiep cua buoc tao moi ban ghi hoac khoi tao mot luong xu ly nay thong qua API, du lieu luu tru hoac trang thai do lop goi phia tren tra ve.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `api/serializers/templates.py`, `accounts.permissions`, `document_templates.models`, `document_templates.utils`, `document_templates.versioning`. Dung cung cap voi cac ham `extract_text_from_docx`, `docx_preview_as_html`, `_replace_in_paragraph` trong module nay.
    Tac dung: Tach rieng trach nhiem tao moi ban ghi hoac khoi tao mot luong xu ly de pham vi tac dong cua `create_docx_from_text` ro rang hon.
    """
    doc = DocxDocument()
    for line in (text or '').split('\n'):
        doc.add_paragraph(line)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

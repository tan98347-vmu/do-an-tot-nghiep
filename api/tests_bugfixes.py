"""Integration tests cho cac fix trong danh_sach_loi_he_thong.md (bug 6, 9, 5, 3)."""

import json
import shutil
import uuid
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from django.conf import settings
from django.contrib.auth.models import User
from django.test import TestCase, override_settings
from django.urls import reverse

from accounts.models import Company, CompanyStatus, CompanyUserMembership
from accounts.storage_paths import company_storage_slug
from ai_engine.models import ChatSession
from documents.models import Document
from document_templates.models import DocumentTemplate


class _FakeLlm:
    def __init__(self, content='OK'):
        self.content = content

    def invoke(self, _messages):
        return SimpleNamespace(content=self.content)


class TemplateDeleteInUseTests(TestCase):
    """Bug 9: chan xoa mau dang duoc su dung, cho phep xoa khi force=true."""

    def setUp(self):
        self.user = User.objects.create_user(username='owner9', password='secret')
        self.client.force_login(self.user)
        self.tmpl = DocumentTemplate.objects.create(
            owner=self.user,
            title='Mau dang dung',
            content='Xin chao {{ ho_ten }}',
            source_type=DocumentTemplate.SOURCE_MANUAL,
            visibility=DocumentTemplate.VISIBILITY_PRIVATE,
        )

    def test_delete_unused_template_succeeds(self):
        resp = self.client.delete(reverse('api:template_detail', args=[self.tmpl.id]))
        self.assertEqual(resp.status_code, 204)
        self.tmpl.refresh_from_db()
        self.assertTrue(self.tmpl.is_deleted)

    def test_delete_in_use_template_blocked_then_forced(self):
        Document.objects.create(
            title='Van ban sinh tu mau',
            owner=self.user,
            template=self.tmpl,
            content='noi dung',
        )
        # Khong co force -> 409
        resp = self.client.delete(reverse('api:template_detail', args=[self.tmpl.id]))
        self.assertEqual(resp.status_code, 409)
        self.assertEqual(resp.json().get('usage_count'), 1)
        self.tmpl.refresh_from_db()
        self.assertFalse(self.tmpl.is_deleted)
        # Co force -> 204
        resp2 = self.client.delete(
            reverse('api:template_detail', args=[self.tmpl.id]) + '?force=true'
        )
        self.assertEqual(resp2.status_code, 204)
        self.tmpl.refresh_from_db()
        self.assertTrue(self.tmpl.is_deleted)


class VariableDetectionTests(TestCase):
    """Bug 6: nhan dien bien co khoang trang `{{ ten }}`."""

    def test_get_variables_with_spaces(self):
        tmpl = DocumentTemplate.objects.create(
            owner=User.objects.create_user(username='owner6', password='x'),
            title='Mau bien',
            content='Kinh gui {{ ho_ten }}, chuc vu {{chuc_vu}}.',
            source_type=DocumentTemplate.SOURCE_MANUAL,
        )
        self.assertEqual(set(tmpl.get_variables()), {'ho_ten', 'chuc_vu'})

    def test_render_with_spaces(self):
        tmpl = DocumentTemplate.objects.create(
            owner=User.objects.create_user(username='owner6b', password='x'),
            title='Mau render',
            content='A {{ ho_ten }} B {{chuc_vu}}',
            source_type=DocumentTemplate.SOURCE_MANUAL,
        )
        self.assertEqual(
            tmpl.render({'ho_ten': 'Nam', 'chuc_vu': 'GD'}),
            'A Nam B GD',
        )


class DocxTemplateVariableDetectionTests(TestCase):
    """Mau nguon DOCX: detect bien tu file .docx ke ca khi `content` bi lech."""

    def test_variables_detected_from_docx_when_content_stale(self):
        import io

        from django.core.files.base import ContentFile
        from docx import Document as WordDocument

        # File DOCX chua bien, mo phong file sau khi sua tren Collabora.
        doc = WordDocument()
        doc.add_paragraph('Hop dong giua {{Ten_Doi_Tac}} va cong ty.')
        doc.add_paragraph('Nguoi dai dien: {{Ten_Nguoi_Dai_Dien}} - {{Chuc_vu}}')
        buf = io.BytesIO()
        doc.save(buf)

        tmpl = DocumentTemplate.objects.create(
            owner=User.objects.create_user(username='owner_docx', password='x'),
            title='Mau DOCX bien',
            # `content` co tinh bi lech (gia tri nhap nhap ban dau), khong co bien.
            content='GGGG',
            source_type=DocumentTemplate.SOURCE_DOCX,
        )
        tmpl.docx_file.save('mau_docx_bien.docx', ContentFile(buf.getvalue()))

        detected = set(tmpl.get_variables())
        self.assertEqual(
            detected,
            {'Ten_Doi_Tac', 'Ten_Nguoi_Dai_Dien', 'Chuc_vu'},
        )


class AiDocCreateCompanyStorageTests(TestCase):
    def setUp(self):
        self.company = Company.objects.create(
            code='vnnet-test',
            slug='vnnet-test',
            name='VNNET Test',
            status=CompanyStatus.ACTIVE,
        )
        self.user = User.objects.create_user(username='vnnet_user', password='secret')
        CompanyUserMembership.objects.create(
            company=self.company,
            user=self.user,
            local_username='vnnet_user',
            role='company_user',
            is_active=True,
        )
        self.template = DocumentTemplate.objects.create(
            owner=self.user,
            company=self.company,
            title='Mau sinh van ban',
            content='Xin chao {{ ho_ten }}',
            source_type=DocumentTemplate.SOURCE_MANUAL,
            visibility=DocumentTemplate.VISIBILITY_PRIVATE,
            status='approved',
        )

    def test_generated_document_file_uses_current_company_storage_prefix(self):
        self.client.force_login(self.user)
        payload = {
            'template_id': self.template.pk,
            'variables': {'ho_ten': 'Nguyen Van A'},
            'doc_title': 'Van ban VNNET',
        }
        media_root = Path(settings.BASE_DIR) / '.codex-tmp-pyc' / f'test-api-ai-doc-{uuid.uuid4().hex}'
        media_root.mkdir(parents=True, exist_ok=True)
        try:
            with override_settings(MEDIA_ROOT=str(media_root)):
                response = self.client.post(
                    reverse('api:ai_doc_create'),
                    data=json.dumps(payload),
                    content_type='application/json',
                )
        finally:
            shutil.rmtree(media_root, ignore_errors=True)

        self.assertEqual(response.status_code, 201)
        document = Document.objects.get(pk=response.json()['id'])
        expected_prefix = f'companies/{company_storage_slug(self.company)}/generated_docs/'
        self.assertTrue(document.output_file.name.startswith(expected_prefix))

        detail_response = self.client.get(reverse('api:document_detail', args=[document.pk]))
        self.assertEqual(detail_response.status_code, 200)


class ChatAiCompanyContextTests(TestCase):
    def setUp(self):
        self.company = Company.objects.create(
            code='chat-company',
            slug='chat-company',
            name='Chat Company',
            status=CompanyStatus.ACTIVE,
            company_context='Ngu canh rieng cua Chat Company.',
        )
        self.user = User.objects.create_user(username='chat_user', password='secret')
        CompanyUserMembership.objects.create(
            company=self.company,
            user=self.user,
            local_username='chat_user',
            role='company_user',
            is_active=True,
        )
        self.client.force_login(self.user)

    def test_chat_endpoint_uses_company_context_without_name_error(self):
        with patch('ai_engine.rag_engine.get_llm', return_value=_FakeLlm('CHAT_OK')):
            response = self.client.post(
                reverse('api:chat_message'),
                data=json.dumps({'q': 'Hoi nhanh'}),
                content_type='application/json',
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()['answer'], 'CHAT_OK')
        session = ChatSession.objects.get(pk=response.json()['session_id'])
        self.assertEqual(session.company_id, self.company.pk)


class NotificationMarkAllTests(TestCase):
    """Bug 5: endpoint danh dau da doc tat ca + badge ve 0 sau khi mark-all."""

    def test_mark_all_endpoint_ok(self):
        user = User.objects.create_user(username='owner5', password='x')
        self.client.force_login(user)
        resp = self.client.post(reverse('api:aggregate_notification_mark_all_read'))
        self.assertEqual(resp.status_code, 200)
        self.assertIn('updated', resp.json())

    def test_mark_all_clears_unread_badge(self):
        from document_templates.models import TemplateReviewNotification

        user = User.objects.create_user(username='owner5b', password='x')
        actor = User.objects.create_user(username='actor5b', password='x')
        self.client.force_login(user)
        tmpl = DocumentTemplate.objects.create(
            owner=user,
            title='Mau thong bao',
            content='abc',
            source_type=DocumentTemplate.SOURCE_MANUAL,
        )
        # 2 thong bao duyet chua doc -> badge = 2
        for _ in range(2):
            TemplateReviewNotification.objects.create(
                recipient=user,
                template=tmpl,
                action=TemplateReviewNotification.ACTION_APPROVE,
                actor=actor,
            )
        count_url = reverse('api:aggregate_notification_unread_count')
        self.assertEqual(self.client.get(count_url).json()['count'], 2)
        # mark-all -> badge ve 0
        self.client.post(reverse('api:aggregate_notification_mark_all_read'))
        self.assertEqual(self.client.get(count_url).json()['count'], 0)

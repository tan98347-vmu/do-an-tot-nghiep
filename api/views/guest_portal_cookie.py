"""
Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
Vai tro backend: File `api/views/guest_portal_cookie.py` giu hoac ho tro luong backend cho chat, RAG, OCR, prefill ho so, sinh van ban, luu session va quan ly tri thuc AI.
Vai tro cua no trong frontend: Cac man `/chat`, `/rag`, `/ai-doc`, `/guest` va cac dialog AI phu tro phu thuoc vao ket qua ma file nay sinh ra.
Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`.
Tac dung: Bao dam prompt, ket qua AI, session hoi thoai, du lieu trich xuat va chi muc RAG phuc vu dung ngu canh cua nguoi dung hien tai.
"""

import json
import re
import secrets
import shutil
import time
import urllib.request
import unicodedata
from pathlib import Path

from django.conf import settings
from django.http import HttpResponse
from django.utils import timezone
from rest_framework import status
from rest_framework.decorators import api_view, authentication_classes, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response

_GUEST_COOKIE_SID = 'guest_sid'
_GUEST_COOKIE_USERNAME = 'guest_username'
_GUEST_META_FILE = 'guest_meta.json'

def _vn_norm(text: str) -> str:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_vn_norm` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    nfkd = unicodedata.normalize('NFKD', text)
    return ''.join(c for c in nfkd if not unicodedata.combining(c)).lower().strip()

def _guest_root() -> Path:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_guest_root` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    root = Path(settings.MEDIA_ROOT) / 'guest_sessions'
    root.mkdir(parents=True, exist_ok=True)
    return root

def _cleanup_stale_guest_dirs(max_age_seconds: int = 60 * 60 * 12) -> None:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_cleanup_stale_guest_dirs` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    root = _guest_root()
    now = time.time()
    for child in root.iterdir():
        try:
            if child.is_dir() and now - child.stat().st_mtime > max_age_seconds:
                shutil.rmtree(child, ignore_errors=True)
        except Exception:
            continue

def _safe_cookie(value: str | None, pattern: str, fallback: str) -> str:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_safe_cookie` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    if value and re.fullmatch(pattern, value):
        return value
    return fallback

def _ensure_guest_identity(request):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_ensure_guest_identity` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    cached = getattr(request, '_guest_identity', None)
    if cached is not None:
        return cached

    _cleanup_stale_guest_dirs()
    sid = _safe_cookie(request.COOKIES.get(_GUEST_COOKIE_SID), r'[a-f0-9]{32}', secrets.token_hex(16))
    username = _safe_cookie(
        request.COOKIES.get(_GUEST_COOKIE_USERNAME),
        r'temp_[a-f0-9]{8}',
        f'temp_{secrets.token_hex(4)}',
    )
    guest_dir = _guest_root() / sid
    guest_dir.mkdir(parents=True, exist_ok=True)
    cached = (sid, username, guest_dir)
    setattr(request, '_guest_identity', cached)
    return cached

def _meta_path(guest_dir: Path) -> Path:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_meta_path` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    return guest_dir / _GUEST_META_FILE

def _read_guest_meta(request) -> dict:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_read_guest_meta` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    _, username, guest_dir = _ensure_guest_identity(request)
    path = _meta_path(guest_dir)
    if not path.exists():
        return {'username': username}
    try:
        data = json.loads(path.read_text(encoding='utf-8'))
        if not isinstance(data, dict):
            return {'username': username}
        data['username'] = username
        return data
    except Exception:
        return {'username': username}

def _write_guest_meta(request, meta: dict) -> dict:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_write_guest_meta` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    _, username, guest_dir = _ensure_guest_identity(request)
    payload = dict(meta)
    payload['username'] = username
    _meta_path(guest_dir).write_text(json.dumps(payload, ensure_ascii=True), encoding='utf-8')
    return payload

def _guest_file(request, meta_key: str) -> Path | None:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_guest_file` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    _, _, guest_dir = _ensure_guest_identity(request)
    meta = _read_guest_meta(request)
    filename = meta.get(meta_key)
    if not filename:
        return None
    path = (guest_dir / filename).resolve()
    try:
        path.relative_to(guest_dir.resolve())
    except Exception:
        return None
    if not path.exists():
        return None
    return path

def _clear_generated_document_meta(request) -> None:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_clear_generated_document_meta` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem xu ly du lieu hoac thao tac lien quan toi van ban truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can xu ly du lieu hoac thao tac lien quan toi van ban nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc xu ly du lieu hoac thao tac lien quan toi van ban de cac endpoint cung file tai su dung dung mot quy tac.
    """
    meta = _read_guest_meta(request)
    path = _guest_file(request, 'document_path')
    if path and path.exists():
        try:
            path.unlink()
        except Exception:
            pass
    for key in ('document_path', 'document_title', 'document_created_at', 'document_updated_at'):
        meta.pop(key, None)
    _write_guest_meta(request, meta)

def _guest_response(request, data, *, http_status=status.HTTP_200_OK) -> Response:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_guest_response` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    sid, username, _ = _ensure_guest_identity(request)
    response = Response(data, status=http_status)
    response.set_cookie(_GUEST_COOKIE_SID, sid, httponly=False, samesite='Lax')
    response.set_cookie(_GUEST_COOKIE_USERNAME, username, httponly=False, samesite='Lax')
    return response

def _attach_guest_cookies(request, response):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_attach_guest_cookies` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    sid, username, _ = _ensure_guest_identity(request)
    response.set_cookie(_GUEST_COOKIE_SID, sid, httponly=False, samesite='Lax')
    response.set_cookie(_GUEST_COOKIE_USERNAME, username, httponly=False, samesite='Lax')
    return response

def _store_guest_template(request, docx_file) -> Path:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_store_guest_template` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem xu ly du lieu hoac thao tac lien quan toi mau van ban truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can xu ly du lieu hoac thao tac lien quan toi mau van ban nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc xu ly du lieu hoac thao tac lien quan toi mau van ban de cac endpoint cung file tai su dung dung mot quy tac.
    """
    _, _, guest_dir = _ensure_guest_identity(request)
    meta = _read_guest_meta(request)

    old_template = _guest_file(request, 'template_path')
    if old_template and old_template.exists():
        try:
            old_template.unlink()
        except Exception:
            pass

    ext = Path(docx_file.name or 'template.docx').suffix or '.docx'
    template_filename = f'template_{secrets.token_hex(4)}{ext}'
    template_path = guest_dir / template_filename
    with template_path.open('wb') as fh:
        for chunk in docx_file.chunks():
            fh.write(chunk)

    meta['template_path'] = template_filename
    meta['template_name'] = docx_file.name or template_filename
    _write_guest_meta(request, meta)
    _clear_generated_document_meta(request)
    return template_path

def _extract_docx_text(doc) -> str:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_extract_docx_text` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem trich xuat noi dung hoac gia tri trung gian truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can trich xuat noi dung hoac gia tri trung gian nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc trich xuat noi dung hoac gia tri trung gian de cac endpoint cung file tai su dung dung mot quy tac.
    """
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
    return '\n'.join(parts)

def _apply_replacements_to_docx(template_path: Path, replacements: dict[str, str]) -> None:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_apply_replacements_to_docx` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    from docx import Document

    doc = Document(str(template_path))

    

    def _apply_to_para(para):
        """
        Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
        Vai tro backend: Ham `_apply_to_para` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
        Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
        Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
        Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
        """
        for orig_text, var_name in replacements.items():
            if not orig_text:
                continue
            placeholder = f'{{{{{var_name}}}}}'
            for run in para.runs:
                if orig_text in run.text:
                    run.text = run.text.replace(orig_text, placeholder)
            full = ''.join(run.text for run in para.runs)
            if orig_text in full and para.runs:
                para.runs[0].text = full.replace(orig_text, placeholder)
                for run in para.runs[1:]:
                    run.text = ''

    for para in doc.paragraphs:
        _apply_to_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _apply_to_para(para)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _apply_to_para(para)
        for para in section.footer.paragraphs:
            _apply_to_para(para)

    doc.save(str(template_path))

def _auto_detect_template_variables(template_path: Path, content: str, existing_vars: list[str]) -> tuple[str, list[str]]:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_auto_detect_template_variables` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem xu ly du lieu hoac thao tac lien quan toi mau van ban truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can xu ly du lieu hoac thao tac lien quan toi mau van ban nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc xu ly du lieu hoac thao tac lien quan toi mau van ban de cac endpoint cung file tai su dung dung mot quy tac.
    """
    from accounts.models import GlobalAIConfig
    from ai_engine.doc_creator import _extract_json_object, _repair_json
    from ai_engine.rag_engine import get_llm
    from langchain_core.messages import HumanMessage, SystemMessage

    cfg = GlobalAIConfig.get_config()
    print(
        f"[guest_parse_template] auto_detect start | file={template_path.name!r} "
        f"| chars={len(content)} | model={cfg.ai_model!r} | ollama={settings.OLLAMA_BASE_URL}"
    )

    ping_started = time.perf_counter()
    try:
        with urllib.request.urlopen(f"{settings.OLLAMA_BASE_URL}/api/tags", timeout=10) as ping_resp:
            ping_ms = (time.perf_counter() - ping_started) * 1000
            print(
                f"[guest_parse_template] ollama ping ok | status={getattr(ping_resp, 'status', '?')} "
                f"| elapsed_ms={ping_ms:.0f}"
            )
    except Exception as ping_exc:
        ping_ms = (time.perf_counter() - ping_started) * 1000
        print(f"[guest_parse_template] ollama ping failed | elapsed_ms={ping_ms:.0f} | error={ping_exc}")

    llm = get_llm()
    system_prompt = (
        "Ban la chuyen gia phan tich mau van ban hanh chinh. "
        "Xac dinh cac thong tin cu the can dien va thay bang placeholder {{ten_bien}} "
        "(snake_case, khong dau). "
        "Tra ve JSON voi 3 truong: "
        "\"content\": noi dung da thay placeholder, "
        "\"variables\": danh sach ten bien, "
        "\"replacements\": mapping tu chuoi goc sang ten bien."
    )
    human_prompt = f"Van ban goc:\n\n{content[:4000]}"

    invoke_started = time.perf_counter()
    resp = llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)])
    invoke_ms = (time.perf_counter() - invoke_started) * 1000
    raw_content = getattr(resp, 'content', '') or ''
    print(
        f"[guest_parse_template] llm response ok | elapsed_ms={invoke_ms:.0f} "
        f"| response_chars={len(raw_content)}"
    )

    parse_started = time.perf_counter()
    result = json.loads(_repair_json(_extract_json_object(str(resp.content).strip())))
    parse_ms = (time.perf_counter() - parse_started) * 1000
    print(f"[guest_parse_template] json parse ok | elapsed_ms={parse_ms:.0f}")

    modified_content = result.get('content', content)
    detected_vars = [str(v) for v in result.get('variables', existing_vars)]
    replacements = result.get('replacements', {})
    print(f"[guest_parse_template] detected vars count={len(detected_vars)} | replacements={len(replacements)}")

    if replacements:
        replace_started = time.perf_counter()
        _apply_replacements_to_docx(template_path, replacements)
        replace_ms = (time.perf_counter() - replace_started) * 1000
        print(f"[guest_parse_template] docx replacement ok | elapsed_ms={replace_ms:.0f}")

    return modified_content, detected_vars

def _guest_document_payload(request):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_guest_document_payload` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem xu ly du lieu hoac thao tac lien quan toi van ban truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can xu ly du lieu hoac thao tac lien quan toi van ban nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc xu ly du lieu hoac thao tac lien quan toi van ban de cac endpoint cung file tai su dung dung mot quy tac.
    """
    meta = _read_guest_meta(request)
    doc_path = _guest_file(request, 'document_path')
    if not doc_path:
        return None

    _, username, _ = _ensure_guest_identity(request)
    created_at = meta.get('document_created_at') or timezone.now().isoformat()
    updated_at = meta.get('document_updated_at') or created_at
    title = meta.get('document_title') or doc_path.stem
    template_name = meta.get('template_name')
    return {
        'id': 0,
        'title': title,
        'content': '',
        'doc_number': '',
        'status': 'final',
        'visibility': 'private',
        'share_status': 'active',
        'owner_id': None,
        'owner_name': username,
        'owner_username': username,
        'group_id': None,
        'template_title': template_name,
        'is_archived': False,
        'has_file': True,
        'created_at': created_at,
        'updated_at': updated_at,
        'source_type': 'guest_ai',
        'notes': 'Van ban tam thoi se bi xoa khi session guest ket thuc.',
        'version_number': 1,
        'version_count': 1,
        'is_favorite': False,
    }

def _guest_html_page(html_body: str) -> str:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_guest_html_page` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
    """
    return f"""<!DOCTYPE html>
<html lang="vi">
<head>
<meta charset="UTF-8">
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    font-family: 'Times New Roman', Times, serif;
    font-size: 14pt;
    background: #e0e0e0;
    padding: 24px;
    line-height: 1.6;
  }}
  .page {{
    background: #ffffff;
    max-width: 210mm;
    min-height: 297mm;
    margin: 0 auto;
    padding: 20mm 25mm 20mm 30mm;
    box-shadow: 0 2px 12px rgba(0,0,0,0.18);
  }}
  h1,h2,h3,h4 {{ font-family: 'Times New Roman', Times, serif; }}
  table {{ border-collapse: collapse; width: 100%; margin: 8px 0; }}
  td, th {{ border: 1px solid #888; padding: 4px 8px; }}
  p {{ margin-bottom: 6px; }}
</style>
</head>
<body>
  <div class="page">{html_body}</div>
</body>
</html>"""

def _extract_pdf_text(pdf_file) -> str:
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `_extract_pdf_text` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem trich xuat noi dung hoac gia tri trung gian truoc khi endpoint chinh phan hoi.
    Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can trich xuat noi dung hoac gia tri trung gian nhung khong nen tu xu ly o client.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
    Tac dung: Co lap rieng buoc trich xuat noi dung hoac gia tri trung gian de cac endpoint cung file tai su dung dung mot quy tac.
    """
    text = ''
    try:
        import pypdf

        reader = pypdf.PdfReader(pdf_file)
        parts = []
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                parts.append(extracted)
        text = '\n'.join(parts)
    except Exception:
        pass

    if not text.strip():
        try:
            pdf_file.seek(0)
            import pdfplumber

            with pdfplumber.open(pdf_file) as pdf:
                text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
        except Exception:
            pass

    return text.strip()

@api_view(['GET'])
@authentication_classes([])
@permission_classes([AllowAny])
def guest_session_info(request):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `guest_session_info` la endpoint hoac diem vao REST cua file `api/views/guest_portal_cookie.py`, chiu trach nhiem quan ly vong doi phien lam viec theo request hien tai.
    Vai tro cua no trong frontend: Frontend goi truc tiep endpoint nay khi nguoi dung can quan ly vong doi phien lam viec tu man hinh tuong ung.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Dung cung cap voi cac ham `_vn_norm`, `_guest_root`, `_cleanup_stale_guest_dirs` trong module nay.
    Tac dung: Bien request thanh JSON, file hoac ma trang thai dung cho thao tac quan ly vong doi phien lam viec tren giao dien.
    """
    _, username, _ = _ensure_guest_identity(request)
    meta = _read_guest_meta(request)
    payload = _guest_document_payload(request)
    return _guest_response(
        request,
        {
            'username': username,
            'template_name': meta.get('template_name'),
            'has_template': _guest_file(request, 'template_path') is not None,
            'has_document': payload is not None,
            'document': payload,
        },
    )

@api_view(['POST'])
@authentication_classes([])
@permission_classes([AllowAny])
def guest_cleanup(request):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `guest_cleanup` la endpoint hoac diem vao REST cua file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai theo request hien tai.
    Vai tro cua no trong frontend: Frontend goi truc tiep endpoint nay khi nguoi dung can thuc hien phan xu ly chuyen trach cua symbol hien tai tu man hinh tuong ung.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Dung cung cap voi cac ham `_vn_norm`, `_guest_root`, `_cleanup_stale_guest_dirs` trong module nay.
    Tac dung: Bien request thanh JSON, file hoac ma trang thai dung cho thao tac thuc hien phan xu ly chuyen trach cua symbol hien tai tren giao dien.
    """
    sid = request.COOKIES.get(_GUEST_COOKIE_SID)
    if sid and re.fullmatch(r'[a-f0-9]{32}', sid):
        guest_dir = _guest_root() / sid
        shutil.rmtree(guest_dir, ignore_errors=True)

    response = Response({'cleared': True})
    response.delete_cookie(_GUEST_COOKIE_SID)
    response.delete_cookie(_GUEST_COOKIE_USERNAME)
    return response

@api_view(['POST'])
@authentication_classes([])
@permission_classes([AllowAny])
def guest_parse_template(request):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `guest_parse_template` la endpoint hoac diem vao REST cua file `api/views/guest_portal_cookie.py`, chiu trach nhiem xu ly du lieu hoac thao tac lien quan toi mau van ban theo request hien tai.
    Vai tro cua no trong frontend: Frontend goi truc tiep endpoint nay khi nguoi dung can xu ly du lieu hoac thao tac lien quan toi mau van ban tu man hinh tuong ung.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Dung cung cap voi cac ham `_vn_norm`, `_guest_root`, `_cleanup_stale_guest_dirs` trong module nay.
    Tac dung: Bien request thanh JSON, file hoac ma trang thai dung cho thao tac xu ly du lieu hoac thao tac lien quan toi mau van ban tren giao dien.
    """
    docx_file = request.FILES.get('docx_file')
    auto_detect = str(request.data.get('auto_detect', 'false')).lower() == 'true'
    if not docx_file:
        return _guest_response(
            request,
            {'detail': 'Can file docx_file.'},
            http_status=status.HTTP_400_BAD_REQUEST,
        )

    try:
        from docx import Document

        template_path = _store_guest_template(request, docx_file)
        doc = Document(str(template_path))
        full_text = _extract_docx_text(doc)
        variables = sorted(set(re.findall(r'\{\{\s*(\w+)\s*\}\}', full_text)))
        preview = full_text[:800]
        auto_detect_applied = False

        if auto_detect and full_text:
            started_at = time.perf_counter()
            try:
                full_text, variables = _auto_detect_template_variables(template_path, full_text, variables)
                preview = full_text[:800]
                auto_detect_applied = True
                total_ms = (time.perf_counter() - started_at) * 1000
                print(f"[guest_parse_template] auto_detect done | total_elapsed_ms={total_ms:.0f}")
            except Exception as e:
                total_ms = (time.perf_counter() - started_at) * 1000
                print(f"[guest_parse_template] auto_detect failed | total_elapsed_ms={total_ms:.0f} | error={e}")

        meta = _read_guest_meta(request)
        return _guest_response(
            request,
            {
                'variables': variables,
                'detected_vars': variables,
                'preview': preview,
                'template_name': meta.get('template_name'),
                'auto_detect': auto_detect_applied,
            },
        )
    except Exception as e:
        return _guest_response(
            request,
            {'detail': f'Loi doc DOCX: {e}'},
            http_status=status.HTTP_400_BAD_REQUEST,
        )

@api_view(['POST'])
@authentication_classes([])
@permission_classes([AllowAny])
def guest_parse_info(request):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `guest_parse_info` la endpoint hoac diem vao REST cua file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai theo request hien tai.
    Vai tro cua no trong frontend: Frontend goi truc tiep endpoint nay khi nguoi dung can thuc hien phan xu ly chuyen trach cua symbol hien tai tu man hinh tuong ung.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Dung cung cap voi cac ham `_vn_norm`, `_guest_root`, `_cleanup_stale_guest_dirs` trong module nay.
    Tac dung: Bien request thanh JSON, file hoac ma trang thai dung cho thao tac thuc hien phan xu ly chuyen trach cua symbol hien tai tren giao dien.
    """
    info_file = request.FILES.get('info_file')
    if not info_file:
        return _guest_response(
            request,
            {'detail': 'Can file info_file.'},
            http_status=status.HTTP_400_BAD_REQUEST,
        )

    try:
        import openpyxl

        wb = openpyxl.load_workbook(info_file, read_only=True, data_only=True)
        ws = wb.active
        values = {}
        header_hints = {'ten bien', 'variable', 'key', 'bien', 'name'}

        for i, row in enumerate(ws.iter_rows(min_row=1)):
            if len(row) < 2:
                continue
            key = str(row[0].value or '').strip()
            val = str(row[1].value or '').strip()
            if not key:
                continue
            if i == 0 and _vn_norm(key) in header_hints:
                continue
            values[key] = val

        wb.close()
        return _guest_response(request, {'values': values})
    except Exception as e:
        return _guest_response(
            request,
            {'detail': f'Loi doc file thong tin: {e}'},
            http_status=status.HTTP_400_BAD_REQUEST,
        )

@api_view(['POST'])
@authentication_classes([])
@permission_classes([AllowAny])
def guest_parse_pdf(request):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `guest_parse_pdf` la endpoint hoac diem vao REST cua file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai theo request hien tai.
    Vai tro cua no trong frontend: Frontend goi truc tiep endpoint nay khi nguoi dung can thuc hien phan xu ly chuyen trach cua symbol hien tai tu man hinh tuong ung.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Dung cung cap voi cac ham `_vn_norm`, `_guest_root`, `_cleanup_stale_guest_dirs` trong module nay.
    Tac dung: Bien request thanh JSON, file hoac ma trang thai dung cho thao tac thuc hien phan xu ly chuyen trach cua symbol hien tai tren giao dien.
    """
    pdf_file = request.FILES.get('pdf_file')
    if not pdf_file:
        return _guest_response(
            request,
            {'detail': 'Can file pdf_file.'},
            http_status=status.HTTP_400_BAD_REQUEST,
        )

    raw_vars = request.data.get('variables', '[]')
    try:
        variables = json.loads(raw_vars) if isinstance(raw_vars, str) else list(raw_vars)
    except Exception:
        variables = []

    pdf_text = _extract_pdf_text(pdf_file)
    if not pdf_text:
        return _guest_response(
            request,
            {'detail': 'Khong doc duoc noi dung PDF.'},
            http_status=status.HTTP_400_BAD_REQUEST,
        )

    preview = pdf_text[:600]
    if not variables:
        return _guest_response(request, {'values': {}, 'matched': 0, 'total': 0, 'raw_preview': preview})

    try:
        from ai_engine.rag_engine import get_llm
        from langchain_core.messages import HumanMessage, SystemMessage

        llm = get_llm()
        vars_list = ', '.join(f'"{var_name}"' for var_name in variables)
        resp = llm.invoke([
            SystemMessage(
                content=(
                    'Ban la chuyen gia trich xuat thong tin tu van ban. '
                    'Doc noi dung PDF va dien gia tri cho cac bien duoc yeu cau. '
                    'Chi dien nhung bien tim thay ro rang trong noi dung. '
                    'Tra ve JSON thuan: {"ten_bien": "gia_tri", ...}.'
                )
            ),
            HumanMessage(content=f'Bien can dien: [{vars_list}]\n\nNoi dung PDF:\n{pdf_text[:4000]}'),
        ])

        raw = str(resp.content).strip()
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        values = json.loads(match.group()) if match else {}
        values = {
            key: str(value).strip()
            for key, value in values.items()
            if key in variables and str(value).strip()
        }
    except Exception:
        values = {}

    return _guest_response(
        request,
        {
            'values': values,
            'matched': len(values),
            'total': len(variables),
            'raw_preview': preview,
        },
    )

@api_view(['POST'])
@authentication_classes([])
@permission_classes([AllowAny])
def guest_generate(request):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `guest_generate` la endpoint hoac diem vao REST cua file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai theo request hien tai.
    Vai tro cua no trong frontend: Frontend goi truc tiep endpoint nay khi nguoi dung can thuc hien phan xu ly chuyen trach cua symbol hien tai tu man hinh tuong ung.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Dung cung cap voi cac ham `_vn_norm`, `_guest_root`, `_cleanup_stale_guest_dirs` trong module nay.
    Tac dung: Bien request thanh JSON, file hoac ma trang thai dung cho thao tac thuc hien phan xu ly chuyen trach cua symbol hien tai tren giao dien.
    """
    _, _, guest_dir = _ensure_guest_identity(request)
    template_path = _guest_file(request, 'template_path')
    if not template_path:
        docx_file = request.FILES.get('docx_file')
        if not docx_file:
            return _guest_response(
                request,
                {'detail': 'Can upload mau van ban truoc khi sinh.'},
                http_status=status.HTTP_400_BAD_REQUEST,
            )
        template_path = _store_guest_template(request, docx_file)

    raw = request.data.get('values', '{}')
    try:
        values = json.loads(raw) if isinstance(raw, str) else dict(raw)
    except Exception:
        values = {}

    try:
        from docx import Document

        doc = Document(str(template_path))

        

        def _fill_para(para, vals):
            """
            Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
            Vai tro backend: Ham `_fill_para` la helper noi bo cua lop API trong file `api/views/guest_portal_cookie.py`, chiu trach nhiem thuc hien phan xu ly chuyen trach cua symbol hien tai truoc khi endpoint chinh phan hoi.
            Vai tro cua no trong frontend: Frontend Flutter khong goi truc tiep helper nay; endpoint cung file dung no khi giao dien can thuc hien phan xu ly chuyen trach cua symbol hien tai nhung khong nen tu xu ly o client.
            Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Thuong duoc cac ham public nhu `guest_session_info`, `guest_cleanup`, `guest_parse_template` goi lai.
            Tac dung: Co lap rieng buoc thuc hien phan xu ly chuyen trach cua symbol hien tai de cac endpoint cung file tai su dung dung mot quy tac.
            """
            for var, val in vals.items():
                placeholder = f'{{{{{var}}}}}'
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(val))
                full = ''.join(run.text for run in para.runs)
                if placeholder in full and para.runs:
                    para.runs[0].text = full.replace(placeholder, str(val))
                    for run in para.runs[1:]:
                        run.text = ''

        for para in doc.paragraphs:
            _fill_para(para, values)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _fill_para(para, values)

        meta = _read_guest_meta(request)
        old_document = _guest_file(request, 'document_path')
        if old_document and old_document.exists():
            try:
                old_document.unlink()
            except Exception:
                pass

        output_filename = f'document_{secrets.token_hex(4)}.docx'
        output_path = guest_dir / output_filename
        doc.save(str(output_path))

        now = timezone.now().isoformat()
        template_name = meta.get('template_name') or template_path.name
        title = request.data.get('title') or Path(template_name).stem
        meta['document_path'] = output_filename
        meta['document_title'] = title
        meta['document_created_at'] = now
        meta['document_updated_at'] = now
        _write_guest_meta(request, meta)

        return _guest_response(
            request,
            {
                'detail': 'ok',
                'document': _guest_document_payload(request),
            },
        )
    except Exception as e:
        return _guest_response(
            request,
            {'detail': f'Loi tao van ban: {e}'},
            http_status=status.HTTP_400_BAD_REQUEST,
        )

@api_view(['GET'])
@authentication_classes([])
@permission_classes([AllowAny])
def guest_document_detail(request):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `guest_document_detail` la endpoint hoac diem vao REST cua file `api/views/guest_portal_cookie.py`, chiu trach nhiem tra du lieu chi tiet cho mot doi tuong cu the theo request hien tai.
    Vai tro cua no trong frontend: Frontend goi truc tiep endpoint nay khi nguoi dung can tra du lieu chi tiet cho mot doi tuong cu the tu man hinh tuong ung.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Dung cung cap voi cac ham `_vn_norm`, `_guest_root`, `_cleanup_stale_guest_dirs` trong module nay.
    Tac dung: Bien request thanh JSON, file hoac ma trang thai dung cho thao tac tra du lieu chi tiet cho mot doi tuong cu the tren giao dien.
    """
    payload = _guest_document_payload(request)
    if not payload:
        return _guest_response(
            request,
            {'detail': 'Guest document not found.'},
            http_status=status.HTTP_404_NOT_FOUND,
        )
    return _guest_response(request, payload)

@api_view(['GET'])
@authentication_classes([])
@permission_classes([AllowAny])
def guest_document_content_html(request):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `guest_document_content_html` la endpoint hoac diem vao REST cua file `api/views/guest_portal_cookie.py`, chiu trach nhiem xu ly du lieu hoac thao tac lien quan toi van ban theo request hien tai.
    Vai tro cua no trong frontend: Frontend goi truc tiep endpoint nay khi nguoi dung can xu ly du lieu hoac thao tac lien quan toi van ban tu man hinh tuong ung.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Dung cung cap voi cac ham `_vn_norm`, `_guest_root`, `_cleanup_stale_guest_dirs` trong module nay.
    Tac dung: Bien request thanh JSON, file hoac ma trang thai dung cho thao tac xu ly du lieu hoac thao tac lien quan toi van ban tren giao dien.
    """
    import mammoth

    doc_path = _guest_file(request, 'document_path')
    if not doc_path:
        return _guest_response(
            request,
            {'detail': 'Guest document not found.'},
            http_status=status.HTTP_404_NOT_FOUND,
        )

    try:
        with doc_path.open('rb') as fh:
            result = mammoth.convert_to_html(fh)
        html_body = result.value
    except Exception as e:
        html_body = f'<p style="color:red">Loi chuyen doi: {e}</p>'

    return _guest_response(request, {'html': _guest_html_page(html_body)})

@api_view(['GET'])
@authentication_classes([])
@permission_classes([AllowAny])
def guest_document_download(request):
    """
    Thuoc chuc nang nao: Tro ly AI, Hoi dap tai lieu, Sinh van ban tu mau, Guest tao van ban va cac luong AI nen.
    Vai tro backend: Ham `guest_document_download` la endpoint hoac diem vao REST cua file `api/views/guest_portal_cookie.py`, chiu trach nhiem tra tep de frontend tai xuong theo request hien tai.
    Vai tro cua no trong frontend: Frontend goi truc tiep endpoint nay khi nguoi dung can tra tep de frontend tai xuong tu man hinh tuong ung.
    Moi lien he voi nhung ham / source khac: Tuong tac truc tiep voi `api/urls.py`, `ai_engine.models`, `ai_engine.rag_engine`, `ai_engine.rag_index`, `ai_engine.doc_creator`, `accounts.models`. Dung cung cap voi cac ham `_vn_norm`, `_guest_root`, `_cleanup_stale_guest_dirs` trong module nay.
    Tac dung: Bien request thanh JSON, file hoac ma trang thai dung cho thao tac tra tep de frontend tai xuong tren giao dien.
    """
    doc_path = _guest_file(request, 'document_path')
    if not doc_path:
        return _guest_response(
            request,
            {'detail': 'Guest document not found.'},
            http_status=status.HTTP_404_NOT_FOUND,
        )

    meta = _read_guest_meta(request)
    with doc_path.open('rb') as fh:
        content = fh.read()

    title = meta.get('document_title') or doc_path.stem
    response = HttpResponse(
        content,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{title}.docx"'
    return _attach_guest_cookies(request, response)

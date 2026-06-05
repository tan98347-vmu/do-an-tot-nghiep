# Chức năng web liên quan: Tạo mẫu văn bản, Mẫu dùng chung, Mẫu phòng ban của tôi, Mẫu riêng của tôi, Mẫu yêu thích và Tất cả mẫu (Admin).
# Vai trò backend trong luồng: Tệp này giữ phần logic backend dùng chung cho danh sách mẫu, chi tiết mẫu, form tạo mẫu, bulk upload và preview biến, để các flow ở các tab Mẫu dùng chung, Mẫu phòng ban, Mẫu riêng, Mẫu yêu thích, màn chi tiết mẫu, form tạo mẫu và bulk upload không phải lặp lại cùng một rule ở nhiều nơi.
# Đầu vào/đầu ra chính: Giữ các helper, cấu hình và rule backend dùng lại ở nhiều flow khác nhau.
# Người dùng sẽ thấy trên web: Người dùng sẽ thấy dữ liệu, badge trạng thái, quyền nút và thông báo trên các chức năng Tạo mẫu văn bản, Mẫu dùng chung, Mẫu phòng ban của tôi, Mẫu riêng của tôi, Mẫu yêu thích và Tất cả mẫu (Admin) thay đổi đúng theo kết quả nghiệp vụ.

import io
import shutil
import uuid
from pathlib import Path
from unittest.mock import patch

from django.conf import settings
from django.contrib.auth.models import User
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document as WordDocument

from document_templates.models import (
    DocumentTemplate,
    TemplateApprovalLog,
    TemplateReviewNotification,
)

# [Web] `TemplateApiTests` gom một cụm xử lý backend dùng chung cho nhóm màn Mẫu văn bản.

class TemplateApiTests(TestCase):
    # [Web] `setUp` là bước nội bộ của class để hỗ trợ dữ liệu hoặc trạng thái mà nhóm màn Mẫu văn bản đang cần.

    def setUp(self):
        self.user = User.objects.create_user(username='template_owner', password='secret')
        self.admin = User.objects.create_superuser(
            username='template_admin',
            password='secret',
            email='admin@example.com',
        )
        self.client.force_login(self.user)

    # [Web] `_make_media_root` là bước nội bộ của class để hỗ trợ dữ liệu hoặc trạng thái mà nhóm màn Mẫu văn bản đang cần.

    def _make_media_root(self):
        media_root = Path(settings.BASE_DIR) / '.codex-tmp-pyc' / f'test-template-{uuid.uuid4().hex}'
        media_root.mkdir(parents=True, exist_ok=True)
        return media_root

    # [Web] `_build_docx_bytes` là bước nội bộ của class để hỗ trợ dữ liệu hoặc trạng thái mà nhóm màn Mẫu văn bản đang cần.

    def _build_docx_bytes(self, *paragraphs):
        doc = WordDocument()
        for paragraph in paragraphs or ('Noi dung mac dinh',):
            doc.add_paragraph(paragraph)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # [Web] `test_create_docx_template_requires_docx_file` là bước nội bộ của class để hỗ trợ dữ liệu hoặc trạng thái mà nhóm màn Mẫu văn bản đang cần.

    def test_create_docx_template_requires_docx_file(self):
        response = self.client.post(
            reverse('api:template_list'),
            data={
                'title': 'Mau DOCX loi',
                'content': '<p>Noi dung</p>',
                'source_type': 'docx',
                'visibility': 'private',
            },
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn('docx_file', response.json())

    # [Web] `test_create_docx_template_with_file_persists_docx_source` là bước nội bộ của class để hỗ trợ dữ liệu hoặc trạng thái mà nhóm màn Mẫu văn bản đang cần.

    def test_create_docx_template_with_file_persists_docx_source(self):
        media_root = self._make_media_root()
        try:
            with override_settings(MEDIA_ROOT=str(media_root)):
                upload = SimpleUploadedFile(
                    'mau-hop-dong.docx',
                    self._build_docx_bytes('Hop dong dich vu {{ho_ten}}'),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                )

                response = self.client.post(
                    reverse('api:template_list'),
                    data={
                        'title': 'Mau hop dong',
                        'content': '<p>Hop dong dich vu {{ho_ten}}</p>',
                        'source_type': 'docx',
                        'visibility': 'private',
                        'docx_file': upload,
                    },
                )

                self.assertEqual(response.status_code, 201, response.content)
                payload = response.json()
                self.assertTrue(payload['has_docx_source'])

                template = DocumentTemplate.objects.get(pk=payload['id'])
                self.assertTrue(bool(template.docx_file))
        finally:
            shutil.rmtree(media_root, ignore_errors=True)

    # [Web] `test_render_as_docx_prefers_docx_source_for_docx_templates` là bước nội bộ của class để hỗ trợ dữ liệu hoặc trạng thái mà nhóm màn Mẫu văn bản đang cần.

    def test_render_as_docx_prefers_docx_source_for_docx_templates(self):
        media_root = self._make_media_root()
        try:
            with override_settings(MEDIA_ROOT=str(media_root)):
                template = DocumentTemplate.objects.create(
                    owner=self.user,
                    title='Mau DOCX uu tien file goc',
                    description='',
                    content='<p>Noi dung editor {{ho_ten}}</p>',
                    source_type=DocumentTemplate.SOURCE_DOCX,
                    status='approved',
                    visibility='private',
                )
                template.docx_file.save(
                    'mau.docx',
                    ContentFile(self._build_docx_bytes('Mau DOCX goc {{ho_ten}}')),
                    save=True,
                )
                template.refresh_from_db()

                rendered = template.render_as_docx({'ho_ten': 'Nguyen Van A'})
                rendered_doc = WordDocument(rendered)
                rendered_text = '\n'.join(
                    paragraph.text for paragraph in rendered_doc.paragraphs if paragraph.text
                )

                self.assertIn('Mau DOCX goc Nguyen Van A', rendered_text)
                self.assertNotIn('Noi dung editor', rendered_text)
        finally:
            shutil.rmtree(media_root, ignore_errors=True)

    # [Web] `test_export_docx_template_without_original_file_returns_conflict` là bước nội bộ của class để hỗ trợ dữ liệu hoặc trạng thái mà nhóm màn Mẫu văn bản đang cần.

    def test_export_docx_template_without_original_file_returns_conflict(self):
        template = DocumentTemplate.objects.create(
            owner=self.user,
            title='Mau DOCX mat file',
            description='',
            content='<p>Noi dung van con</p>',
            source_type=DocumentTemplate.SOURCE_DOCX,
            status='approved',
            visibility='private',
        )

        response = self.client.get(reverse('api:template_export', args=[template.pk]))

        self.assertEqual(response.status_code, 409, response.content)
        self.assertEqual(response.json()['code'], 'no_docx_source')

    # [Web] `test_preview_docx_template_without_original_file_returns_conflict` là bước nội bộ của class để hỗ trợ dữ liệu hoặc trạng thái mà nhóm màn Mẫu văn bản đang cần.

    def test_preview_docx_template_without_original_file_returns_conflict(self):
        template = DocumentTemplate.objects.create(
            owner=self.user,
            title='Mau DOCX mat file preview',
            description='',
            content='<p>Noi dung van con</p>',
            source_type=DocumentTemplate.SOURCE_DOCX,
            status='approved',
            visibility='private',
        )

        response = self.client.get(reverse('api:template_preview_pdf', args=[template.pk]))

        self.assertEqual(response.status_code, 409, response.content)
        self.assertEqual(response.json()['code'], 'no_docx_source')

    # [Web] `test_import_from_url_docx_returns_source_docx_payload` là bước nội bộ của class để hỗ trợ dữ liệu hoặc trạng thái mà nhóm màn Mẫu văn bản đang cần.

    @patch('api.views.templates._fetch_remote_template_source')
    def test_import_from_url_docx_returns_source_docx_payload(self, fetch_mock):
        docx_bytes = self._build_docx_bytes('Tai lieu tu URL')
        fetch_mock.return_value = {
            'content': 'Tai lieu tu URL',
            'title': 'Tai lieu tu URL',
            'source_kind': 'docx',
            'resolved_url': 'https://example.com/files/tai-lieu.docx',
            'raw_bytes': docx_bytes,
        }

        response = self.client.post(
            reverse('api:template_import_from_url'),
            data={
                'source_url': 'https://example.com/files/tai-lieu.docx',
                'auto_detect': 'false',
            },
        )

        self.assertEqual(response.status_code, 200, response.content)
        payload = response.json()
        self.assertIn('source_docx', payload)
        self.assertEqual(payload['source_filename'], 'tai-lieu.docx')

    # [Web] `test_template_approve_creates_review_notification` là bước nội bộ của class để hỗ trợ dữ liệu hoặc trạng thái mà nhóm màn Mẫu văn bản đang cần.

    def test_template_approve_creates_review_notification(self):
        template = DocumentTemplate.objects.create(
            owner=self.user,
            title='Mau cho duyet',
            description='',
            content='<p>Noi dung</p>',
            source_type=DocumentTemplate.SOURCE_MANUAL,
            status='pending',
            visibility='public',
        )

        self.client.force_login(self.admin)
        response = self.client.post(
            reverse('api:template_approve', args=[template.pk]),
            data={'comment': 'Da duyet'},
        )

        self.assertEqual(response.status_code, 200, response.content)
        template.refresh_from_db()
        self.assertEqual(template.status, 'approved')
        self.assertTrue(
            TemplateApprovalLog.objects.filter(
                template=template,
                action=TemplateApprovalLog.ACTION_APPROVE,
                actor=self.admin,
            ).exists()
        )
        notification = TemplateReviewNotification.objects.get(template=template, recipient=self.user)
        self.assertEqual(notification.action, TemplateApprovalLog.ACTION_APPROVE)
        self.assertEqual(notification.comment, 'Da duyet')

    # [Web] `test_template_reject_creates_review_notification_and_keeps_template` là bước nội bộ của class để hỗ trợ dữ liệu hoặc trạng thái mà nhóm màn Mẫu văn bản đang cần.

    def test_template_reject_creates_review_notification_and_keeps_template(self):
        template = DocumentTemplate.objects.create(
            owner=self.user,
            title='Mau bi tu choi',
            description='',
            content='<p>Noi dung</p>',
            source_type=DocumentTemplate.SOURCE_MANUAL,
            status='pending',
            visibility='public',
        )

        self.client.force_login(self.admin)
        response = self.client.post(
            reverse('api:template_reject', args=[template.pk]),
            data={'comment': 'Can bo sung noi dung'},
        )

        self.assertEqual(response.status_code, 200, response.content)
        template.refresh_from_db()
        self.assertEqual(template.status, 'rejected')
        self.assertEqual(template.visibility, DocumentTemplate.VISIBILITY_PRIVATE)
        self.assertTrue(DocumentTemplate.objects.filter(pk=template.pk).exists())
        notification = TemplateReviewNotification.objects.get(template=template, recipient=self.user)
        self.assertEqual(notification.action, TemplateApprovalLog.ACTION_REJECT)
        self.assertEqual(notification.comment, 'Can bo sung noi dung')

    def test_template_reject_requires_reason(self):
        template = DocumentTemplate.objects.create(
            owner=self.user,
            title='Mau can ly do tu choi',
            description='',
            content='<p>Noi dung</p>',
            source_type=DocumentTemplate.SOURCE_MANUAL,
            status='pending',
            visibility='public',
        )

        self.client.force_login(self.admin)
        response = self.client.post(reverse('api:template_reject', args=[template.pk]), data={})

        self.assertEqual(response.status_code, 400, response.content)
        self.assertEqual(response.json()['detail'], 'Phai co ly do tu choi mau.')
